import os
from pathlib import Path
from datetime import datetime


def open_and_read_word(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return f"📝 **محتوى الملف:**\n\n{text[:3000]}"
    except Exception as e:
        return f"❌ فشل قراءة الملف: {e}"


def append_to_word(path: str, text: str) -> str:
    try:
        from docx import Document
        if os.path.exists(path):
            doc = Document(path)
        else:
            doc = Document()
        doc.add_paragraph(text)
        doc.save(path)
        return f"✅ تم إضافة النص إلى: {Path(path).name}"
    except Exception as e:
        return f"❌ فشل التعديل: {e}"


def create_word_document(path: str, title: str = "", content: str = "") -> str:
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        if title:
            heading = doc.add_heading(title, level=1)
        if content:
            doc.add_paragraph(content)
        doc.add_paragraph(f"\nتم الإنشاء: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.save(path)
        return f"✅ تم إنشاء الملف: {Path(path).name}"
    except Exception as e:
        return f"❌ فشل الإنشاء: {e}"


def replace_text_in_word(path: str, old_text: str, new_text: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        replaced = 0
        for para in doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replaced += 1
        doc.save(path)
        return f"✅ تم استبدال {replaced} نص في: {Path(path).name}"
    except Exception as e:
        return f"❌ فشل الاستبدال: {e}"


def add_table_to_word(path: str, headers: list, rows: list) -> str:
    try:
        from docx import Document
        if os.path.exists(path):
            doc = Document(path)
        else:
            doc = Document()
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            header_cells[i].text = h
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
        doc.save(path)
        return f"✅ تم إضافة الجدول إلى: {Path(path).name}"
    except Exception as e:
        return f"❌ فشل إضافة الجدول: {e}"
